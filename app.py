from flask import Flask, render_template, request, send_file
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches
import os

from docx.oxml import OxmlElement
from docx.oxml.ns import qn


# =========================
# APP INIT
# =========================
app = Flask(__name__)

# Create outputs folder
if not os.path.exists("outputs"):
    os.makedirs("outputs")


# =========================
# REMOVE TABLE BORDERS
# =========================
def remove_table_borders(table):
    tbl = table._element
    tblPr = tbl.tblPr

    borders = OxmlElement('w:tblBorders')

    for edge in ('top','left','bottom','right','insideH','insideV'):
        tag = OxmlElement(f'w:{edge}')
        tag.set(qn('w:val'), 'nil')
        borders.append(tag)

    tblPr.append(borders)


# =========================
# ADD HORIZONTAL LINE (FIXED)
# =========================
def add_horizontal_line(doc):
    p = doc.add_paragraph()
    p_format = p._element.get_or_add_pPr()

    border = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')

    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '12')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '000000')

    border.append(bottom)
    p_format.append(border)


# =========================
# HOME PAGE
# =========================
@app.route('/')
def index():
    return render_template('index.html')


# =========================
# WORD GENERATION
# =========================
@app.route('/generate_word', methods=['POST'])
def generate_word():

    data = request.form
    file_path = "outputs/question_paper.docx"

    doc = Document()

    # =========================
    # FONT SETTING
    # =========================
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    # =========================
    # HEADER TABLE
    # =========================
    table = doc.add_table(rows=5, cols=3)

    table.cell(0,0).merge(table.cell(0,2))
    table.cell(1,0).merge(table.cell(1,2))
    table.cell(2,1).merge(table.cell(3,1))
    table.cell(4,0).merge(table.cell(4,1))

    # School
    p = table.cell(0,0).paragraphs[0]
    p.text = data.get('school','')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Exam
    p = table.cell(1,0).paragraphs[0]
    p.text = data.get('exam','')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Subject
    p = table.cell(2,1).paragraphs[0]
    p.text = data.get('subject','')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Left side
    table.cell(2,0).text = f"Max. Marks: {data.get('marks','')}"

    # Right side
    table.cell(2,2).text = f"Class: {data.get('class','')}"
    table.cell(3,2).text = f"Duration: {data.get('duration','')}"

    # Student details
    table.cell(4,0).text = "Name of Student:\n\n"
    table.cell(4,2).text = "Roll No:\n\nDate:\n"

    # Remove borders
    remove_table_borders(table)

    doc.add_paragraph("\n")


    # =========================
    # INSTRUCTIONS
    # =========================
    p = doc.add_paragraph()
    run = p.add_run("General Instructions")
    run.bold = True
    run.font.size = Pt(14)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    instructions = data.get('instructions','').split("\n")

    for line in instructions:
        if line.strip():
            p = doc.add_paragraph(line)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1

    # ✅ LINE AFTER INSTRUCTIONS (NOW WORKING)
    add_horizontal_line(doc)

    doc.add_paragraph("\n")


    # =========================
    # QUESTION FUNCTION
    # =========================
    def add_questions(section_name, prefix):

        # Section Title
        p = doc.add_paragraph()
        run = p.add_run(section_name)
        run.bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph()

        i = 1

        while True:

            q_text = data.get(f"{prefix}_q_{i}")

            if not q_text:
                break

            marks = data.get(f"{prefix}_marks_{i}")
            q_type = data.get(f"{prefix}_type_{i}")

            # Question text
            p = doc.add_paragraph(f"{i}. {q_text}")
            p.paragraph_format.space_after = Pt(6)

            # Marks (right side)
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p.add_run(f"{marks}")

            # Image support
            image = request.files.get(f"{prefix}_img_{i}")
            if image and image.filename != "":
                image_path = os.path.join("outputs", image.filename)
                image.save(image_path)
                doc.add_picture(image_path, width=Inches(5))

            # MCQ options
            if q_type == "mcq":

                options = [
                    data.get(f"{prefix}_a_{i}"),
                    data.get(f"{prefix}_b_{i}"),
                    data.get(f"{prefix}_c_{i}"),
                    data.get(f"{prefix}_d_{i}")
                ]

                labels = ['a)', 'b)', 'c)', 'd)']

                for idx, opt in enumerate(options):
                    if opt:
                        doc.add_paragraph(f"{labels[idx]} {opt}")

            doc.add_paragraph()
            i += 1


    # =========================
    # CALL SECTIONS
    # =========================
    add_questions("Section A", "secA")
    add_questions("Section B", "secB")
    add_questions("Section C", "secC")


    # =========================
    # SAVE
    # =========================
    doc.save(file_path)

    return send_file(file_path, as_attachment=True)


# =========================
# PREVIEW
# =========================
@app.route('/preview', methods=['GET','POST'])
def preview():
    data = request.form
    return render_template("preview.html", data=data)


# =========================
# RUN SERVER
# =========================
if __name__ == '__main__':
    app.run(host='0.0.0.0', port=5000)